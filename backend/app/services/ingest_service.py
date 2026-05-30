from __future__ import annotations
import asyncio
import io
import re
from pathlib import Path
from typing import List, Optional

from ..core.config import get_settings


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text(filename: str, file_bytes: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(file_bytes)
    elif suffix == ".docx":
        return _extract_docx(file_bytes)
    elif suffix in (".txt", ".md"):
        raw = file_bytes.decode("utf-8", errors="replace")
        return _preprocess_structured_txt(raw)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf(data: bytes) -> str:
    """
    Structure-aware PDF text extraction.

    Strategy per page:
      1. Detect tables with PyMuPDF's geometric analyser (find_tables).
      2. For table cells: forward-fill merged cells column-by-column, detect
         single-value "section header" rows (e.g. "IRIFM"), and prefix every
         data row with its section so each row is self-contained for RAG.
      3. For non-table text: emit regular text blocks that don't overlap a
         detected table bounding box.
      4. If table detection fails (rare), fall back to plain text extraction.
    """
    import fitz  # PyMuPDF
    doc = fitz.open(stream=data, filetype="pdf")
    parts: list[str] = []

    for page in doc:
        # ── 1. Detect tables ───────────────────────────────────────────────
        try:
            tab_finder = page.find_tables()
            tables_on_page = tab_finder.tables if tab_finder else []
        except Exception:
            tables_on_page = []

        table_bboxes = [t.bbox for t in tables_on_page]

        # ── 2. Non-table text blocks ───────────────────────────────────────
        for block in page.get_text("blocks"):
            x0, y0, x1, y1 = block[0], block[1], block[2], block[3]
            text = block[4].strip()
            if not text:
                continue
            in_table = any(
                x0 < tx1 and x1 > tx0 and y0 < ty1 and y1 > ty0
                for (tx0, ty0, tx1, ty1) in table_bboxes
            )
            if not in_table:
                parts.append(text)

        # ── 3. Tables with structure preservation ─────────────────────────
        for table in tables_on_page:
            try:
                rows = table.extract()
            except Exception:
                continue
            if not rows:
                continue

            n_cols = max((len(r) for r in rows), default=0)
            if n_cols == 0:
                continue

            # Tokens that identify a column-header row in directory tables.
            _HEADER_TOKENS = {
                "name", "designation", "rly", "bsnl", "phone", "office",
                "residence", "mobile", "fax", "email", "stn", "stdn",
            }

            def _is_col_header_row(row: list) -> bool:
                """True if the row looks like a column-label row (≥2 header tokens)."""
                tokens = [str(c or "").strip().lower() for c in row if c]
                return sum(1 for t in tokens if t in _HEADER_TOKENS) >= 2

            # Pass 1: identify section-header rows (merged single-cell rows) and
            # column-header rows from the ORIGINAL data (before forward-fill).
            is_section_row: list[bool] = []
            is_col_header_row: list[bool] = []
            for row in rows:
                padded = list(row) + [None] * (n_cols - len(row))
                first_val = str(padded[0] or "").strip()
                rest_all_none = all(c is None for c in padded[1:])
                is_section_row.append(bool(first_val and rest_all_none))
                is_col_header_row.append(_is_col_header_row(padded))

            # Pass 2: forward-fill None values for data rows (vertically merged cells)
            prev_vals: list[str] = [""] * n_cols
            filled: list[list[str]] = []
            for row in rows:
                padded = list(row) + [None] * (n_cols - len(row))
                new_row: list[str] = []
                for i, cell in enumerate(padded):
                    if cell is not None:
                        val = " ".join(str(cell).split())
                        prev_vals[i] = val
                        new_row.append(val)
                    else:
                        new_row.append(prev_vals[i])
                filled.append(new_row)

            # Collect the most-recent column headers seen in this table so we
            # can label each data row's values (e.g. "Rly: 83400 | BSNL: 27057490")
            col_headers: list[str] = []
            for idx, row in enumerate(filled):
                if is_col_header_row[idx]:
                    col_headers = [c.strip() for c in row]

            def _label_row(row: list[str]) -> str:
                """
                If column headers are known, emit "Header: value" pairs for
                non-trivial headers (skip dept/name-like ones already in section).
                Otherwise fall back to plain pipe-join.
                """
                if not col_headers:
                    return " | ".join(c for c in row if c.strip())
                pairs = []
                for header, value in zip(col_headers, row):
                    if not value.strip() or value.strip() == "-":
                        continue
                    h = header.strip()
                    # Skip generic positional headers that add no meaning
                    if h.lower() in ("", "designation", "name"):
                        pairs.append(value.strip())
                    else:
                        pairs.append(f"{h}: {value.strip()}")
                return " | ".join(pairs)

            # Pass 3: emit rows with section context
            # Each section is separated by a blank entry so the chunker
            # treats sections as paragraph boundaries (splits on \n\n first).
            current_section = ""
            prev_emitted_section = ""

            for idx, row in enumerate(filled):
                non_empty = [c for c in row if c.strip()]
                if not non_empty:
                    continue

                if is_section_row[idx]:
                    current_section = non_empty[0]
                    continue

                if is_col_header_row[idx]:
                    continue  # headers are embedded into row labels; don't emit as data

                row_text = _label_row(row)
                if not row_text.strip():
                    continue

                if current_section and row[0].strip() != current_section:
                    formatted = f"{current_section} | {row_text}"
                    active_section = current_section
                else:
                    formatted = row_text
                    active_section = row[0].strip()

                # Insert blank-line separator when the section label changes
                # so each section becomes its own chunk boundary
                if prev_emitted_section and active_section != prev_emitted_section:
                    parts.append("")   # blank line → \n\n when joined → chunk split

                parts.append(formatted)
                prev_emitted_section = active_section

        # ── 4. Fallback: no tables detected → plain text ──────────────────
        if not tables_on_page:
            plain = page.get_text("text").strip()
            if plain:
                parts.append(plain)

    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(data))
    parts: List[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # Tables — row by row, cell by cell
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _preprocess_structured_txt(text: str) -> str:
    """
    For structured TXT / MD files (directories, personnel lists, timelines)
    reconstruct each "entry" as a self-contained line that includes its section
    heading as a prefix.

    Pattern recognised:
      - A block separated by blank lines that contains a date (DD-MM-YYYY or
        similar) is treated as a **person/event entry** and flushed as a single
        pipe-separated line prefixed by the current section heading.
      - A block with no date that looks like a section heading (contains known
        keywords or is a single short line) becomes the running context.
      - For plain narrative text the algorithm is a safe no-op: no date blocks
        → entries accumulate and are joined normally, headings pass through.

    Example output:
      Previous Director Generals of IRIFM: Reena Ranjan | Director General | 02-11-2020 to 31-10-2021
    """
    IMAGE_EXT = re.compile(r'\.(jpg|jpeg|png|gif|webp|svg|bmp)$', re.I)
    DATE_RE   = re.compile(r'\b\d{2}[/\-]\d{2}[/\-]\d{2,4}\b')

    # Keywords that mark section-heading blocks.
    # Must be *strong* qualifiers — avoid generic role words like "Director"
    # which also appear as designations inside entries.
    HEADING_KEYWORDS = [
        'previous', 'present', 'past', 'former', 'current',
        'heads of', 'directors at',
        'irifm', 'campus-ii', 'campus ii',
        'railway', 'institution',
    ]

    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    raw_blocks = [b.strip() for b in text.split('\n\n') if b.strip()]

    # Pre-scan: is this file "structured" (many short lines + dates)?
    # If not, return the original text unchanged (safe for prose documents).
    short_line_count = sum(1 for b in raw_blocks if len(b) < 40)
    has_dates        = any(DATE_RE.search(b) for b in raw_blocks)
    is_structured    = has_dates and short_line_count > len(raw_blocks) * 0.4

    if not is_structured:
        return text

    result: List[str] = []
    current_heading   = ""
    entry_parts: List[str] = []

    def _flush_entry():
        if not entry_parts:
            return
        entry_str = " | ".join(entry_parts)
        result.append(f"{current_heading}: {entry_str}" if current_heading else entry_str)
        entry_parts.clear()

    for block in raw_blocks:
        lines = [l.strip() for l in block.split('\n') if l.strip()]
        # Drop image file references
        lines = [l for l in lines if not IMAGE_EXT.search(l)]
        if not lines:
            continue

        block_text = ' '.join(lines)

        # Decide if this block is a section heading
        is_heading = any(kw in block_text.lower() for kw in HEADING_KEYWORDS)

        if is_heading:
            _flush_entry()
            current_heading = block_text
            # Don't emit the heading itself — it will be prepended to each entry
            continue

        # Accumulate lines into the current entry
        entry_parts.extend(lines)

        # A date signals the end of an entry — flush immediately
        if DATE_RE.search(block_text):
            _flush_entry()

    _flush_entry()   # flush anything remaining

    # Fall back to original text if we produced nothing useful
    return '\n\n'.join(result) if result else text


# ---------------------------------------------------------------------------
# Recursive character splitter
# ---------------------------------------------------------------------------

SEPARATORS = ["\n\n", "\n", ". ", " ", ""]


def _split_text(text: str, chunk_size: int, chunk_overlap: int, separators: List[str]) -> List[str]:
    """
    Recursively split text using a hierarchy of separators.
    Produces chunks of at most `chunk_size` characters.
    Adjacent chunks share `chunk_overlap` characters at their boundary
    (end of chunk N == start of chunk N+1's overlap window).
    """
    if not text.strip():
        return []

    sep = separators[0] if separators else ""
    remaining = separators[1:] if len(separators) > 1 else []

    splits = text.split(sep) if sep else list(text)

    good_chunks: List[str] = []
    current = ""

    for part in splits:
        joined = (current + sep + part) if current else part
        if len(joined) <= chunk_size:
            current = joined
        else:
            if current.strip():
                good_chunks.append(current)
            if len(part) > chunk_size and remaining:
                good_chunks.extend(_split_text(part, chunk_size, chunk_overlap, remaining))
                current = ""
            else:
                current = part

    if current.strip():
        good_chunks.append(current)

    # Merge very small fragments into their neighbour
    merged: List[str] = []
    for chunk in good_chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        if merged and len(merged[-1]) + len(chunk) + 1 <= chunk_size:
            merged[-1] = merged[-1] + sep + chunk
        else:
            merged.append(chunk)

    if chunk_overlap <= 0 or len(merged) <= 1:
        return merged

    # Correct sliding-window overlap: each chunk starts with the last
    # `chunk_overlap` chars of the PREVIOUS chunk, capped to chunk_size.
    result: List[str] = [merged[0]]
    for i in range(1, len(merged)):
        tail = merged[i - 1][-chunk_overlap:]
        candidate = tail + " " + merged[i]
        # Trim to chunk_size if the overlap pushes us over
        result.append(candidate[:chunk_size] if len(candidate) > chunk_size else candidate)
    return result


def chunk_text(text: str) -> list[str]:
    settings = get_settings()
    raw = _split_text(text, settings.chunk_size, settings.chunk_overlap, SEPARATORS)
    # Filter extremely short scraps (page numbers, lone punctuation, etc.)
    # Lowered from 80 → 20 so structured table rows aren't discarded
    return [c for c in raw if len(c.strip()) > 20]


# ---------------------------------------------------------------------------
# Agentic chunking (LLM-based boundary detection)
# ---------------------------------------------------------------------------

async def _agentic_chunk(
    text: str,
    progress_queue: Optional[asyncio.Queue] = None,
) -> list[str]:
    """
    LLM-based agentic chunking.

    Splits the text into paragraphs then asks Ollama for each paragraph:
    "Does this continue the current chunk, or start a new topic?"
    Chunks stay together until Ollama says NEW or until the hard ceiling
    `max_agentic_chunk_chars` is reached.

    Falls back transparently to the character splitter if Ollama is unreachable.

    If `progress_queue` is provided, emits chunking progress events into it so
    callers (e.g. an SSE endpoint) can relay live updates to the client without
    the connection going idle during the long CPU-bound LLM loop.
    """
    from . import ollama_client as _ollama

    settings = get_settings()
    max_chars = settings.max_agentic_chunk_chars

    # Split on paragraph boundaries
    paragraphs = [p.strip() for p in re.split(r"\n\n+", text) if p.strip()]
    if not paragraphs:
        return []

    total = len(paragraphs)
    chunks: list[str] = []
    current_parts: list[str] = []
    current_len = 0

    def _flush() -> None:
        if current_parts:
            chunks.append("\n\n".join(current_parts))
        current_parts.clear()
        nonlocal current_len
        current_len = 0

    def _enqueue(para: str) -> None:
        """
        Add a paragraph to the current accumulator.  If the paragraph itself
        exceeds the hard ceiling (e.g. an un-broken wall of PDF text), split
        it immediately with the character splitter so it never becomes an
        un-embeddable chunk.
        """
        nonlocal current_len
        if len(para) > max_chars:
            # Single paragraph is already too large — split it right now
            _flush()
            chunks.extend(chunk_text(para))
        else:
            current_parts.append(para)
            current_len += len(para) + 2

    for idx, para in enumerate(paragraphs):
        # Emit progress so the SSE stream stays alive between Ollama calls
        if progress_queue is not None:
            await progress_queue.put({
                "type": "progress",
                "phase": "chunking",
                "current": idx + 1,
                "total": total,
            })

        # Bootstrap: first paragraph always starts the first chunk
        if not current_parts and current_len == 0:
            _enqueue(para)
            continue

        # Hard ceiling — flush before this paragraph pushes us over
        if current_len + len(para) + 2 > max_chars:
            _flush()
            _enqueue(para)
            continue

        # Ask Ollama: boundary or continuation?
        current_text = "\n\n".join(current_parts)
        try:
            is_new = await _ollama.classify_chunk_boundary(current_text, para)
        except Exception:
            is_new = False  # safe default: keep accumulating

        if is_new:
            _flush()
            _enqueue(para)
        else:
            current_parts.append(para)
            current_len += len(para) + 2  # +2 for the \n\n separator

    # Flush the last in-progress chunk
    _flush()

    result = [c for c in chunks if len(c.strip()) > 20]

    # Safety: if agentic chunking produced nothing, fall back to char splitter
    if not result:
        return chunk_text(text)

    return result


# ---------------------------------------------------------------------------
# Public entry point (async — uses agentic chunking)
# ---------------------------------------------------------------------------

async def extract_and_chunk(
    filename: str,
    file_bytes: bytes,
    progress_queue: Optional[asyncio.Queue] = None,
) -> list[str]:
    """
    Extract text and chunk using LLM-based agentic boundary detection.

    TXT / MD path:
      1. Try structured preprocessing (directory / personnel lists with dates).
         If the file is recognised as structured, each entry becomes its own
         self-contained chunk — no further processing needed.
      2. Otherwise apply agentic chunking so narrative text is split at topic
         boundaries rather than character limits.

    PDF / DOCX path:
      Extract text (table-aware for PDFs) then apply agentic chunking.

    If `progress_queue` is provided it is forwarded to `_agentic_chunk`, which
    emits per-paragraph progress events so the SSE stream stays alive during
    the long chunking phase.
    """
    suffix = Path(filename).suffix.lower()

    if suffix in (".txt", ".md"):
        raw = file_bytes.decode("utf-8", errors="replace")
        processed = _preprocess_structured_txt(raw)
        if processed != raw:
            # Structured directory/timeline — each \n\n entry is its own chunk
            entries = [e.strip() for e in processed.split("\n\n") if len(e.strip()) > 20]
            if entries:
                return entries
        # Narrative TXT/MD — agentic chunking
        text = re.sub(r"\r\n", "\n", raw)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return await _agentic_chunk(text, progress_queue)

    # PDF / DOCX
    text = extract_text(filename, file_bytes)
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return await _agentic_chunk(text, progress_queue)
